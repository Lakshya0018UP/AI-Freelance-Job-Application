from app import db ,app

from flask import request, jsonify
from models import Jobs,User,Applied
from flask_jwt_extended import create_access_token,get_jwt_identity,jwt_required,create_refresh_token
from datetime import datetime 
import os
import fitz
import docx
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

#Get all Jobs

@app.route("/api/jobs",methods=['GET'])
def get_jobs():
    jobs=Jobs.query.all()
    result=[job.to_json() for job in jobs]
    return jsonify(result)


#create a job
@app.route("/api/jobs",methods=['POST'])
@jwt_required()
def create_jobs():
    try:
        identity=get_jwt_identity()
        print("Here is the identity",identity)
        print(identity["role"])
        if identity["role"] not in ("admin","professional"):
            return jsonify({"msg":"Job can only be created by the admin or Professional"})
        
        data=request.get_json()
        required_fields=['title','budget','description']
        for field in required_fields:
            if field not in data:
                return{"msg":f'missing required field:{field}'}

        title=data.get("title")
        description=data.get("description")
        budget=data.get("budget")
        skills=data.get("skills")

        new_job=Jobs(title=title,description=description,budget=budget,skills=skills,user_id=identity["id"])
        db.session.add(new_job)
        db.session.commit()
        return jsonify({"msg":"Job added succesfully"}),201
    except Exception as e:
        db.session.rollback()
        return jsonify({"error":str(e)}),500

#Delete a Job
@app.route("/api/delete_job/<int:id>",methods=['DELETE'])
@jwt_required()
def delete_job(id):
    try:
        identity=get_jwt_identity()
        if identity['role'] not in ['admin','professional']:
            return jsonify ({"msg":"You are not authorized to delete this Job"}),401

        #job_to_be_deleted=Job.query.get_or_404(id) It gives that if record is not found, it automatically gives an error, hence making the specific error reduntant of use 
        job_to_be_deleted=Jobs.query.get(id)
        if job_to_be_deleted is None:
            return jsonify ({"msg":"The Paticular Job has not been found"}),404

        db.session.delete(job_to_be_deleted)
        db.session.commit()

        #jsonify({"msg":f'The said record has been delete{job_to_be_deleted}'})
        return jsonify({"msg":"Job Deleted"})
    except Exception as e:
        db.session.rollback()
        return jsonify ({"error":str(e)}),500
    

#update Job
@app.route("/api/update_jobs/<int:id>",methods=['PATCH'])
def update_jobs(id):
    try:
        job_to_be_update=Jobs.query.get(id)
        if job_to_be_update is None:
            return jsonify({"msg":"No record found of such Job"})
        data=request.get_json()

        job_to_be_update.title=data.get("title",job_to_be_update.title)
        job_to_be_update.description=data.get("description",job_to_be_update.description)
        job_to_be_update.budget=data.get("budget",job_to_be_update.budget)

        db.session.commit()
        return jsonify ({"msg":"The data has been updated"})
    except Exception as e:
        db.session.rollback()
        return jsonify ({"error":str(e)})
    
@app.route("/api/users",methods=['GET'])
def users():
    Users=User.query.all()
    result=[user.to_json() for user in Users]
    return jsonify(result)

@app.route('/api/signup',methods=['POST'])
def signup():
    data=request.get_json()

    username=data.get('username')
    name=data.get('name')
    email=data.get('email')
    password_hash=data.get('password_hash')
    role=data.get('role')

    existing_user=User.query.filter_by(email=email).first()
    if existing_user:
        return jsonify({"msg":"The user already exists"})
    
    required_fields=['username','name','email','password_hash','role']
    for field in required_fields:
        if field not in data:
            return jsonify({"error":f'A field is missing:{field}'})

    new_user=User(username=username,name=name,email=email,role=role)
    new_user.set_password(password_hash)

    db.session.add(new_user)
    db.session.commit()

    return jsonify({"msg":"The User has been added"})


@app.route("/api/login",methods=['POST'])
def login():
    data=request.get_json()
    username=data.get('username')
    password=data.get('password')
    
    user_in=User.query.filter_by(username=username).first()
    if not user_in or not user_in.check_password(password):
        return jsonify ({"msg":"Either password or username incorrect"})
    
    access_token=create_access_token(identity={"id":user_in.id,"role":user_in.role,"username":user_in.username,"email":user_in.email,"name":user_in.name})
    refresh_token=create_refresh_token(identity={"id":user_in.id,"role":user_in.role,"username":user_in.username,"email":user_in.email,"name":user_in.name})
    return jsonify ({"msg":"Hurray!!You are logged in","access_token":access_token,"refresh_token":refresh_token}),200

@app.route("/api/view_job/<int:id>",methods=['POST'])
@jwt_required()
def view_job(id):
    job=Jobs.query.get(id)
    if job is None:
        return jsonify({"msg":"The Job is not found"})
    return jsonify(job.to_json())

@app.route("/api/apply/<int:id>",methods=['POST'])
@jwt_required()
def applied(id):
    applied_job=Jobs.query.get(id)
    if applied_job is None:
        return jsonify({"msg":"The Job does not exist"}),404
    
    user_in=get_jwt_identity()
    user_already=Applied.query.filter_by(job_id=id,username=user_in['username']).first()
    if user_already:
        return jsonify({"error":"You have already applied for this role"}),404
    
    data=request.get_json()
    proposal=data.get("proposal")
    bid_amount=data.get("bid_amount")
    user_applied=Applied(username=user_in['username'],name=user_in['name'],email=user_in['email'],job_id=id,proposal=proposal,bid_amount=bid_amount,freelancer_id=user_in['id'])
    db.session.add(user_applied)
    db.session.commit()
    return jsonify({"msg":"You have applied for the Job"}),201

@app.route("/api/applied_jobs/<int:id>",methods=['GET'])
@jwt_required()
def applied_jobs(id):
    user_in=get_jwt_identity()
    if user_in['role'] not in ["admin","professional"]:
        return jsonify({"msg":"You are not authorized to check this out"}),404
    all_applies=Applied.query.filter_by(job_id=id).all()
    if all_applies is None:
        return jsonify({"msg":"No one has applied for this job"})
    result=[apply.to_json() for apply in all_applies]
    return jsonify(result)

#TODO : Add the functionality update the status of the job
@app.route("/api/update_status/<int:id>",methods=['PATCH'])
@jwt_required()
def update_status(id):
    user_in=get_jwt_identity()
    if user_in['role'] not in ['admin','professional']:
        return jsonify({"msg":"You are not authotized for this page"}),401
    
    applied_change=Applied.query.get(id)
    if applied_change is None:
        return jsonify({"msg":"The paticular job doesnt exist"}),404
    
    data=request.get_json()
    print(data)

    applied_change.status=data.get("status",applied_change.status)
    db.session.commit()
    return jsonify({"msg":"The status has been updated"}),200

@app.route("/api/view_profile",methods=['GET'])
@jwt_required()
def view_profile():
    user_in=get_jwt_identity()
    if user_in['role']=='freelancer':
        proposals=Applied.query.filter_by(freelancer_id=user_in['id']).all()
        if proposals is None:
            return jsonify({"msg":"You have not applied for any job"}),404
        result=[apply.to_json() for apply in proposals]
    elif user_in['role']=='professional':
        jobs=Jobs.query.filter_by(user_id=user_in['id']).all()
        if jobs is None:
            return jsonify({"msg":"You have not posted any job"}),404
        result=[job.to_json() for job in jobs]
    return jsonify(user_in,result),200

@app.route("/api/refresh",methods=['GET'])
@jwt_required(refresh=True)
def refresh():
    identity=get_jwt_identity()
    new_access_token=create_access_token(identity=identity)

    return jsonify({"access_token":new_access_token}),200


# # ------------ Resume Utilities ------------
# ideal_keywords = [
#     "python", "flask", "django", "rest api", "sql",
#     "html", "css", "javascript", "data scraping", "automation",'java','nodejs',
#     "aws", "azure", "docker", "kubernetes", "git","ci/cd",'github'
# ]

# def extract_text(file_path):
#     if file_path.endswith('.pdf'):
#         return extract_text_from_pdf(file_path)
#     elif file_path.endswith('.docx'):
#         return extract_text_from_docx(file_path)
#     elif file_path.endswith('.txt'):
#         with open(file_path, 'r', encoding='utf-8') as f:
#             return f.read()
#     return ""

# def extract_text_from_pdf(file_path):
#     text = ""
#     with fitz.open(file_path) as doc:
#         for page in doc:
#             text += page.get_text()
#     return text

# def extract_text_from_docx(file_path):
#     doc = docx.Document(file_path)
#     return "\n".join([para.text for para in doc.paragraphs])

# def calculate_resume_score(resume, keywords):
#     resume = resume.lower()
#     matched_keywords = [kw for kw in keywords if kw in resume]
#     score = (len(matched_keywords) / len(keywords)) * 100
#     return round(score, 2), matched_keywords

# def recommend_jobs(resume, jobs, top_n=3):
#     corpus = [resume] + [job.description for job in jobs]
#     vectorizer = TfidfVectorizer(stop_words='english')
#     tfidf_matrix = vectorizer.fit_transform(corpus)

#     similarities = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:]).flatten()
#     job_scores = sorted(zip(jobs, similarities), key=lambda x: x[1], reverse=True)
    
#     return [{
#         "title": job.title,
#         "description": job.description,
#         "similarity_score": round(score * 100, 2)
#     } for job, score in job_scores[:top_n]]


# @app.route('/api/upload_resume', methods=['POST'])
# @jwt_required()
# def upload_resume():
#         if 'resume' not in request.files:
#             return jsonify({"error": "No file uploaded"}), 400

#         file = request.files['resume']
#         filename = file.filename

#         if not filename.lower().endswith(('.pdf', '.docx', '.txt')):
#             return jsonify({"error": "Unsupported file type"}), 400

#         file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
#         file.save(file_path)

#         resume_text = extract_text(file_path)
#         if not resume_text.strip():
#             return jsonify({"error": "Could not extract resume text"}), 500

#         jobs = Jobs.query.all()
#         score, matched_keywords = calculate_resume_score(resume_text, ideal_keywords)
#         recommendations = recommend_jobs(resume_text, jobs)

#         return jsonify({
#             "resume_score": score,
#             "matched_keywords": matched_keywords,
#             "recommended_jobs": recommendations
#         })

ideal_keywords = [
    "python", "flask", "django", "rest api", "sql",
    "html", "css", "javascript", "data scraping", "express.js", "java", "nodejs",
    "aws", "azure", "docker", "kubernetes", "git", "ci/cd", "github"
]
def extract_text(file_path):
    if file_path.endswith('.pdf'):
        return extract_text_from_pdf(file_path)
    elif file_path.endswith('.docx'):
        return extract_text_from_docx(file_path)
    elif file_path.endswith('.txt'):
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    return ""

def extract_text_from_pdf(file_path):
    text = ""
    with fitz.open(file_path) as doc:
        for page in doc:
            text += page.get_text()
    return text

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def calculate_resume_score(resume, keywords):
    resume = resume.lower()
    matched_keywords = [kw for kw in keywords if kw in resume]
    score = (len(matched_keywords) / len(keywords)) * 100
    return round(score, 2), matched_keywords

def recommend_jobs(resume, jobs, top_n=3):
    corpus = [resume] + [job.description for job in jobs]
    vectorizer = TfidfVectorizer(stop_words='english')
    tfidf_matrix = vectorizer.fit_transform(corpus)

    similarities = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:]).flatten()
    job_scores = sorted(zip(jobs, similarities), key=lambda x: x[1], reverse=True)

    return [{
        "title": job.title,
        "description": job.description,
        "similarity_score": round(score * 100, 2)
    } for job, score in job_scores[:top_n]]

@app.route('/api/upload_resume', methods=['POST'])
@jwt_required()
def upload_resume():
    # Look for resume file in any form key
    file = None
    for f in request.files.values():
        file = f
        break

    if file is None:
        return jsonify({"error": "No file uploaded. Please upload your resume file using a form field named 'resume'."}), 400

    filename = file.filename

    if not filename.lower().endswith(('.pdf', '.docx', '.txt')):
        return jsonify({"error": "Unsupported file type. Only PDF, DOCX, and TXT are allowed."}), 400

    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(file_path)

    resume_text = extract_text(file_path)
    if not resume_text.strip():
        return jsonify({"error": "Could not extract any text from the uploaded resume."}), 500

    jobs = Jobs.query.all()
    score, matched_keywords = calculate_resume_score(resume_text, ideal_keywords)
    recommendations = recommend_jobs(resume_text, jobs)

    return jsonify({
        "message": "Resume uploaded and processed successfully.",
        "resume_score": score,
        "matched_keywords": matched_keywords,
        "recommended_jobs": recommendations
    })